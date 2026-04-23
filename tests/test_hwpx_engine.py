from __future__ import annotations

import json
import subprocess
import sys
import zipfile
from pathlib import Path

import pytest

try:
    from docx import Document
except ModuleNotFoundError:  # pragma: no cover - environment dependent
    Document = None

from shawn_hwp.converters.hwpx_engine import (
    extract_hwpx_text,
    parse_hwpx_to_model,
    run_docx_to_hwpx_conversion,
    run_hwpx_to_docx_conversion,
    run_hwpx_to_md_conversion,
    run_md_to_hwpx_conversion,
)
from shawn_hwp.io_hwpx import render_hwpx_section_xml
from shawn_hwp.parse_markdown import parse_markdown


REPO_ROOT = Path(__file__).resolve().parents[1]


def _make_fake_hwpx(path: Path) -> None:
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr(
            "Contents/section0.xml",
            """
            <hp:section xmlns:hp=\"http://www.hancom.co.kr/hwpml/2011/paragraph\">
              <hp:p style-name=\"Title\"><hp:run><hp:t>Sample Title</hp:t></hp:run></hp:p>
              <hp:p><hp:run><hp:t>First paragraph.</hp:t></hp:run></hp:p>
              <hp:p><hp:run><hp:t>Second paragraph.</hp:t></hp:run></hp:p>
            </hp:section>
            """.strip(),
        )


def _make_fake_hwpx_with_inline_runs(path: Path) -> None:
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr(
            "Contents/section0.xml",
            """
            <hp:section xmlns:hp=\"http://www.hancom.co.kr/hwpml/2011/paragraph\">
              <hp:p><hp:run><hp:t>1. 연구개발과제의 필요성</hp:t></hp:run></hp:p>
              <hp:p>
                <hp:run charPrIDRef=\"bold\"><hp:t>Bold</hp:t></hp:run>
                <hp:run><hp:t> plain</hp:t></hp:run>
              </hp:p>
            </hp:section>
            """.strip(),
        )


def _make_fake_hwpx_with_table(path: Path) -> None:
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr(
            "Contents/section0.xml",
            """
            <hp:section xmlns:hp=\"http://www.hancom.co.kr/hwpml/2011/paragraph\">
              <hp:tbl>
                <hp:tr>
                  <hp:tc><hp:p><hp:run><hp:t>Col1</hp:t></hp:run></hp:p></hp:tc>
                  <hp:tc><hp:p><hp:run><hp:t>Col2</hp:t></hp:run></hp:p></hp:tc>
                </hp:tr>
                <hp:tr>
                  <hp:tc><hp:p><hp:run><hp:t>A</hp:t></hp:run></hp:p></hp:tc>
                  <hp:tc><hp:p><hp:run><hp:t>B</hp:t></hp:run></hp:p></hp:tc>
                </hp:tr>
              </hp:tbl>
            </hp:section>
            """.strip(),
        )


def _make_bridge_noise_hwpx(path: Path) -> None:
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr(
            "Contents/section0.xml",
            """
            <hp:section xmlns:hp=\"http://www.hancom.co.kr/hwpml/2011/paragraph\">
              <hp:p><hp:run><hp:t>1-</hp:t></hp:run></hp:p>
              <hp:p><hp:run><hp:t>작성요령(제출 시 삭제)</hp:t></hp:run></hp:p>
              <hp:p><hp:run><hp:t>건국대학교 첨단재생과학연구원 목 차 중장기 육성 계획</hp:t></hp:run></hp:p>
              <hp:p><hp:run><hp:t>실제 본문 문단</hp:t></hp:run></hp:p>
            </hp:section>
            """.strip(),
        )


def _make_bridge_table_echo_hwpx(path: Path) -> None:
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr(
            "Contents/section0.xml",
            """
            <hp:section xmlns:hp=\"http://www.hancom.co.kr/hwpml/2011/paragraph\">
              <hp:tbl>
                <hp:tr>
                  <hp:tc><hp:p><hp:run><hp:t>항목</hp:t></hp:run></hp:p></hp:tc>
                  <hp:tc><hp:p><hp:run><hp:t>내용</hp:t></hp:run></hp:p></hp:tc>
                </hp:tr>
                <hp:tr>
                  <hp:tc><hp:p><hp:run><hp:t>A</hp:t></hp:run></hp:p></hp:tc>
                  <hp:tc><hp:p><hp:run><hp:t>B</hp:t></hp:run></hp:p></hp:tc>
                </hp:tr>
              </hp:tbl>
              <hp:p><hp:run><hp:t>항목 내용 A B</hp:t></hp:run></hp:p>
            </hp:section>
            """.strip(),
        )


def _make_front_layout_table_hwpx(path: Path) -> None:
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr(
            "Contents/section0.xml",
            """
            <hp:section xmlns:hp=\"http://www.hancom.co.kr/hwpml/2011/paragraph\">
              <hp:tbl>
                <hp:tr>
                  <hp:tc><hp:p><hp:run><hp:t>2025년도 이공분야 학술연구지원사업글로컬랩 연구계획서</hp:t></hp:run></hp:p></hp:tc>
                </hp:tr>
              </hp:tbl>
              <hp:p><hp:run><hp:t>1. 연구개발과제의 필요성</hp:t></hp:run></hp:p>
              <hp:tbl>
                <hp:tr>
                  <hp:tc><hp:p><hp:run><hp:t>항목</hp:t></hp:run></hp:p></hp:tc>
                  <hp:tc><hp:p><hp:run><hp:t>내용</hp:t></hp:run></hp:p></hp:tc>
                </hp:tr>
                <hp:tr>
                  <hp:tc><hp:p><hp:run><hp:t>A</hp:t></hp:run></hp:p></hp:tc>
                  <hp:tc><hp:p><hp:run><hp:t>B</hp:t></hp:run></hp:p></hp:tc>
                </hp:tr>
              </hp:tbl>
            </hp:section>
            """.strip(),
        )


def test_parse_hwpx_to_model_reads_blocks(tmp_path: Path):
    source = tmp_path / "sample.hwpx"
    _make_fake_hwpx(source)

    model = parse_hwpx_to_model(source)

    assert model.blocks[0].kind == "heading"
    assert model.blocks[0].text == "Sample Title"
    assert model.blocks[1].kind == "paragraph"
    assert model.blocks[1].text == "First paragraph."


def test_extract_hwpx_text_reads_paragraphs(tmp_path: Path):
    source = tmp_path / "sample.hwpx"
    _make_fake_hwpx(source)

    text = extract_hwpx_text(source)

    assert "# Sample Title" in text
    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_parse_hwpx_to_model_preserves_inline_runs(tmp_path: Path):
    source = tmp_path / "inline.hwpx"
    _make_fake_hwpx_with_inline_runs(source)

    model = parse_hwpx_to_model(source)

    assert model.blocks[1].text == "Bold plain"
    assert len(model.blocks[1].runs) == 2
    assert model.blocks[1].runs[0].text == "Bold"
    assert model.blocks[1].runs[0].bold is True
    assert model.blocks[1].runs[1].text == " plain"


def test_run_hwpx_to_md_conversion_writes_markdown(tmp_path: Path):
    source = tmp_path / "sample.hwpx"
    output = tmp_path / "out.md"
    _make_fake_hwpx(source)

    result = run_hwpx_to_md_conversion(source, output, "hwpx", "md")

    assert output.exists()
    assert result.route == "hwpx-to-md"
    body = output.read_text(encoding="utf-8")
    assert "# Sample Title" in body
    assert "First paragraph." in body


def test_markdown_parser_and_hwpx_writer_roundtrip():
    model = parse_markdown("# Sample Title\n\nFirst paragraph.\n\n| A | B |\n| --- | --- |\n| 1 | 2 |\n")
    xml = render_hwpx_section_xml(model)

    assert "style-name=\"Title\"" in xml
    assert "<hp:tbl>" in xml
    assert "Sample Title" in xml


def test_render_hwpx_section_xml_preserves_multiple_runs():
    model = parse_markdown("First paragraph.\n")
    model.blocks[0].runs = []
    from shawn_hwp.model import InlineRun

    model.blocks[0].runs = [InlineRun(text="Bold", bold=True), InlineRun(text=" plain")]
    model.blocks[0].text = "Bold plain"
    xml = render_hwpx_section_xml(model)

    assert xml.count("<hp:run") >= 2
    assert 'charPrIDRef="bold"' in xml
    assert "Bold" in xml and " plain" in xml


def test_run_md_to_hwpx_conversion_writes_hwpx(tmp_path: Path):
    source = tmp_path / "source.md"
    output = tmp_path / "out.hwpx"
    source.write_text("# Sample Title\n\nFirst paragraph.\n\nSecond paragraph.\n", encoding="utf-8")

    result = run_md_to_hwpx_conversion(source, output, "md", "hwpx")

    assert output.exists()
    assert result.route == "md-to-hwpx"
    extracted = extract_hwpx_text(output)
    assert "# Sample Title" in extracted
    assert "First paragraph." in extracted


def test_extract_hwpx_text_reads_table_as_markdown(tmp_path: Path):
    source = tmp_path / "sample_table.hwpx"
    _make_fake_hwpx_with_table(source)

    text = extract_hwpx_text(source)

    assert "| Col1 | Col2 |" in text
    assert "| --- | --- |" in text
    assert "| A | B |" in text


def test_parse_hwpx_to_model_filters_bridge_noise_and_splits_front_matter(tmp_path: Path):
    source = tmp_path / "bridge_noise.hwpx"
    _make_bridge_noise_hwpx(source)

    model = parse_hwpx_to_model(source)
    texts = [block.text for block in model.blocks if block.kind in {"heading", "paragraph"}]

    assert "1-" not in texts
    assert "작성요령(제출 시 삭제)" not in texts
    assert "건국대학교 첨단재생과학연구원" not in texts
    assert "목 차" in texts
    assert any(block.kind == "heading" and block.text == "목 차" for block in model.blocks)
    assert "실제 본문 문단" not in texts


def test_parse_hwpx_to_model_deduplicates_table_echo_paragraph(tmp_path: Path):
    source = tmp_path / "bridge_echo.hwpx"
    _make_bridge_table_echo_hwpx(source)

    model = parse_hwpx_to_model(source)
    table_blocks = [block for block in model.blocks if block.kind == "table"]
    paragraph_texts = [block.text for block in model.blocks if block.kind == "paragraph"]

    assert len(table_blocks) == 1
    assert "항목 내용 A B" not in paragraph_texts


def test_parse_hwpx_to_model_skips_front_layout_table_but_keeps_body_table(tmp_path: Path):
    source = tmp_path / "front_layout.hwpx"
    _make_front_layout_table_hwpx(source)

    model = parse_hwpx_to_model(source)
    table_blocks = [block for block in model.blocks if block.kind == "table"]
    heading_texts = [block.text for block in model.blocks if block.kind == "heading"]

    assert len(table_blocks) == 1
    assert table_blocks[0].rows[0] == ["항목", "내용"]
    assert "1. 연구개발과제의 필요성" in heading_texts


def test_run_md_table_to_hwpx_roundtrip(tmp_path: Path):
    source = tmp_path / "table.md"
    output = tmp_path / "table.hwpx"
    source.write_text("| Col1 | Col2 |\n| --- | --- |\n| A | B |\n", encoding="utf-8")

    run_md_to_hwpx_conversion(source, output, "md", "hwpx")
    extracted = extract_hwpx_text(output)

    assert "| Col1 | Col2 |" in extracted
    assert "| A | B |" in extracted


@pytest.mark.skipif(Document is None, reason="python-docx not available in test interpreter")
def test_run_docx_to_hwpx_conversion_writes_hwpx(tmp_path: Path):
    source = tmp_path / "source.docx"
    output = tmp_path / "out.hwpx"
    doc = Document()
    doc.add_heading("Sample Title", level=1)
    doc.add_paragraph("First paragraph.")
    doc.save(source)

    result = run_docx_to_hwpx_conversion(source, output, "docx", "hwpx")

    assert output.exists()
    assert result.route == "docx-to-hwpx"
    extracted = extract_hwpx_text(output)
    assert "# Sample Title" in extracted
    assert "First paragraph." in extracted


@pytest.mark.skipif(Document is None, reason="python-docx not available in test interpreter")
def test_run_hwpx_to_docx_conversion_writes_docx(tmp_path: Path):
    source = tmp_path / "sample.hwpx"
    output = tmp_path / "out.docx"
    _make_fake_hwpx(source)

    result = run_hwpx_to_docx_conversion(source, output, "hwpx", "docx")

    assert output.exists()
    assert result.route == "hwpx-to-docx"
    doc = Document(output)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Sample Title" in text
    assert "First paragraph." in text


def test_convert_cli_uses_hwpx_native_engine_for_md(tmp_path: Path):
    source = tmp_path / "sample.hwpx"
    output = tmp_path / "out.md"
    metadata = tmp_path / "meta.json"
    _make_fake_hwpx(source)

    cmd = [
        sys.executable,
        str(REPO_ROOT / "scripts" / "convert.py"),
        "--input",
        str(source),
        "--from",
        "hwpx",
        "--to",
        "md",
        "--output",
        str(output),
        "--emit-metadata",
        str(metadata),
    ]

    result = subprocess.run(cmd, cwd=REPO_ROOT, capture_output=True, text=True, check=True)

    assert "engine=hwpx-native" in result.stdout
    assert output.exists()
    payload = json.loads(metadata.read_text(encoding="utf-8"))
    assert payload["source_format"] == "hwpx"
    assert payload["target_format"] == "md"


@pytest.mark.skipif(Document is None, reason="python-docx not available in test interpreter")
def test_convert_cli_uses_hwpx_native_engine_for_docx(tmp_path: Path):
    source = tmp_path / "sample.hwpx"
    output = tmp_path / "out.docx"
    metadata = tmp_path / "meta.json"
    _make_fake_hwpx(source)

    cmd = [
        sys.executable,
        str(REPO_ROOT / "scripts" / "convert.py"),
        "--input",
        str(source),
        "--from",
        "hwpx",
        "--to",
        "docx",
        "--output",
        str(output),
        "--emit-metadata",
        str(metadata),
    ]

    result = subprocess.run(cmd, cwd=REPO_ROOT, capture_output=True, text=True, check=True)

    assert "engine=hwpx-native" in result.stdout
    assert output.exists()
    payload = json.loads(metadata.read_text(encoding="utf-8"))
    assert payload["source_format"] == "hwpx"
    assert payload["target_format"] == "docx"


def test_convert_cli_uses_hwpx_native_engine_for_md_to_hwpx(tmp_path: Path):
    source = tmp_path / "source.md"
    output = tmp_path / "out.hwpx"
    metadata = tmp_path / "meta.json"
    source.write_text("# Sample Title\n\nFirst paragraph.\n", encoding="utf-8")

    cmd = [
        sys.executable,
        str(REPO_ROOT / "scripts" / "convert.py"),
        "--input",
        str(source),
        "--from",
        "md",
        "--to",
        "hwpx",
        "--output",
        str(output),
        "--emit-metadata",
        str(metadata),
    ]

    result = subprocess.run(cmd, cwd=REPO_ROOT, capture_output=True, text=True, check=True)

    assert "engine=hwpx-native" in result.stdout
    assert output.exists()
    payload = json.loads(metadata.read_text(encoding="utf-8"))
    assert payload["source_format"] == "md"
    assert payload["target_format"] == "hwpx"
    assert "Sample Title" in extract_hwpx_text(output)


@pytest.mark.skipif(Document is None, reason="python-docx not available in test interpreter")
def test_convert_cli_uses_hwpx_native_engine_for_docx_to_hwpx(tmp_path: Path):
    source = tmp_path / "source.docx"
    output = tmp_path / "out.hwpx"
    metadata = tmp_path / "meta.json"
    doc = Document()
    doc.add_heading("Sample Title", level=1)
    doc.add_paragraph("First paragraph.")
    doc.save(source)

    cmd = [
        sys.executable,
        str(REPO_ROOT / "scripts" / "convert.py"),
        "--input",
        str(source),
        "--from",
        "docx",
        "--to",
        "hwpx",
        "--output",
        str(output),
        "--emit-metadata",
        str(metadata),
    ]

    result = subprocess.run(cmd, cwd=REPO_ROOT, capture_output=True, text=True, check=True)

    assert "engine=hwpx-native" in result.stdout
    assert output.exists()
    payload = json.loads(metadata.read_text(encoding="utf-8"))
    assert payload["source_format"] == "docx"
    assert payload["target_format"] == "hwpx"
    assert "Sample Title" in extract_hwpx_text(output)
