"""Native HWPX extraction helpers for SHawn-hwp."""

from __future__ import annotations

import html
import re
import zipfile
from pathlib import Path
import xml.etree.ElementTree as ET
from xml.sax.saxutils import escape
from typing import Literal

from shawn_hwp.converters.stub import ConversionResult
from shawn_hwp.converters.strategy_router import choose_route
from shawn_hwp.io_docx import write_docx
from shawn_hwp.io_hwpx import write_hwpx
from shawn_hwp.io_markdown import render_markdown
from shawn_hwp.model import DocumentModel, InlineRun
from shawn_hwp.parse_markdown import parse_markdown


_TEXT_TAG_CANDIDATES = {
    "t",
    "text",
    "script",
    "p",
    "hp:t",
    "hp:text",
    "hp:script",
    "hp:p",
}

_PARAGRAPH_TAG_CANDIDATES = {
    "p",
    "paragraph",
    "hp:p",
}

_HEADING_NAME_HINTS = {"title", "heading", "header", "chapter", "section"}
_TABLE_ROW_HINTS = {"tr", "row"}
_TABLE_CELL_HINTS = {"td", "tc", "cell"}
_BRIDGE_NOISE_PATTERNS = (
    r"^\d+-$",
    r"^작성요령\(제출 시 삭제\)",
    r"^작성요령",
    r"^\s*작성요령\s*\(제출 시 삭제\)",
    r"^※",
    r"^표에서 파란색으로 작성된 부분은 예시입니다",
    r"^※ 성과 목표 계획표",
    r"^◎ 대상사업\s*:",
)
_FRONT_MATTER_SPLIT_PATTERNS = (
    r"(?=목\s*차)",
    r"(?=건국대학교)",
    r"(?=중장기)",
)
_BODY_START_PATTERNS = (
    r"^1\.\s*연구개발과제의 필요성$",
    r"^연구개발과제의 필요성$",
)
_FRONT_COVER_NOISE_PATTERNS = (
    r"^건국대학교 첨단재생과학연구원$",
    r"^설립$",
    r"^중장기 연구목표를 중심으로 작성",
    r"^중장기 육성 계획에 대하여 자유롭게 기술",
    r"^가\. 연구개발성과의 활용방안$",
    r"^나\. 연구개발성과의 기대효과$",
)
_SECTION_HEADING_PATTERNS = (
    r"^\d+\.\s*.+$",
    r"^\(\d+\)\s*.+$",
)
_ADMIN_TAIL_PATTERNS = (
    r"^\(단위:\s*천원\)",
    r"연구개발기관명",
    r"연구개발비",
    r"현금",
    r"현물",
    r"총계",
    r"간접비 활용계획",
    r"배분계획",
)


def hwpx_available() -> bool:
    return True


def hwpx_docx_available() -> bool:
    try:
        import docx  # noqa: F401
    except ModuleNotFoundError:
        return False
    return True


def _local_name(tag: str) -> str:
    if "}" in tag:
        return tag.rsplit("}", 1)[1]
    return tag


def _iter_xml_members(zf: zipfile.ZipFile) -> list[str]:
    names = [name for name in zf.namelist() if name.lower().endswith(".xml")]
    body_first = sorted(names, key=lambda name: (0 if "section" in name.lower() or "contents" in name.lower() else 1, name))
    return body_first


def _node_text(node: ET.Element) -> str:
    parts: list[str] = []
    if node.text and node.text.strip():
        parts.append(node.text)
    for child in node:
        child_name = _local_name(child.tag).lower()
        if child_name in {"linebreak", "br"}:
            parts.append("\n")
        else:
            child_text = _node_text(child)
            if child_text.strip():
                parts.append(child_text)
        if child.tail and child.tail.strip():
            parts.append(child.tail)
    text = "".join(parts)
    text = html.unescape(text)
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r" ?\n ?", "\n", text)
    return text.strip()


def _extract_table_rows(node: ET.Element) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in node.iter():
        row_name = _local_name(row.tag).lower()
        if row_name not in _TABLE_ROW_HINTS:
            continue
        cells: list[str] = []
        for child in row:
            cell_name = _local_name(child.tag).lower()
            if cell_name in _TABLE_CELL_HINTS:
                cell_text = _node_text(child).strip()
                cells.append(cell_text or " ")
        if cells:
            rows.append(cells)
    return rows


def _style_flags(style_hint: str) -> tuple[bool, bool, bool]:
    hint = style_hint.lower()
    return (
        any(token in hint for token in ("bold", "strong", "fwbold", "weight=700")),
        any(token in hint for token in ("italic", "emphasis", "oblique")),
        any(token in hint for token in ("underline", "under")),
    )


def _extract_inline_runs(node: ET.Element, trace: str) -> list[InlineRun]:
    runs: list[InlineRun] = []
    for child in node:
        child_name = _local_name(child.tag).lower()
        attrs = {(_local_name(k).lower()): v for k, v in child.attrib.items()}
        style_hint = " ".join(attrs.values())
        if child_name == "run":
            raw_text = html.unescape("".join(child.itertext()))
            raw_text = re.sub(r"[\t\r\f\v]+", " ", raw_text)
            raw_text = raw_text.replace("\u00a0", " ")
            if not raw_text.strip():
                continue
            bold, italic, underline = _style_flags(style_hint)
            runs.append(
                InlineRun(
                    text=raw_text,
                    bold=bold,
                    italic=italic,
                    underline=underline,
                    style_hint=style_hint or None,
                    source_trace=f"{trace}:run",
                )
            )
            continue
        if child_name in {"linebreak", "br"}:
            runs.append(InlineRun(text="\n", source_trace=f"{trace}:br"))

    if runs:
        return runs

    text = _node_text(node)
    return [InlineRun(text=text, source_trace=f"{trace}:text")] if text else []


def _normalize_text(text: str) -> str:
    text = html.unescape(text or "")
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r" ?\n ?", "\n", text)
    return text.strip()


def _is_bridge_noise(text: str) -> bool:
    normalized = _normalize_text(text)
    if not normalized:
        return True
    return any(re.search(pattern, normalized) for pattern in _BRIDGE_NOISE_PATTERNS)


def _split_glued_paragraph(text: str) -> list[str]:
    normalized = _normalize_text(text)
    if not normalized:
        return []

    parts = [normalized]
    for pattern in _FRONT_MATTER_SPLIT_PATTERNS:
        next_parts: list[str] = []
        for part in parts:
            split = [chunk.strip() for chunk in re.split(pattern, part) if chunk.strip()]
            next_parts.extend(split or [part])
        parts = next_parts

    return parts


def _should_promote_to_heading(text: str, style_hint: str, body_started: bool = True) -> bool:
    normalized = _normalize_text(text)
    if not normalized:
        return False
    if any(hint in style_hint for hint in _HEADING_NAME_HINTS):
        return True
    if normalized in {"목 차"}:
        return True
    if any(re.match(pattern, normalized) for pattern in _BODY_START_PATTERNS):
        return True
    if re.match(r"^\d+\.\s+.+$", normalized):
        return True
    if body_started and re.match(r"^\(?\d+\)?\.?(\s|$)", normalized):
        return True
    if body_started and len(normalized) <= 80 and re.match(r"^(제?\d+[장절항]?|[IVX]+\.|[0-9]+\.|[가-힣A-Z][가-힣A-Z0-9\s·()\-/]+)$", normalized):
        return True
    return False


def _table_signature(rows: list[list[str]]) -> tuple[tuple[str, ...], ...]:
    return tuple(tuple(_normalize_text(cell) for cell in row) for row in rows)


def _detect_parser_state(text: str, current: Literal["front", "toc", "body", "admin-tail"]) -> Literal["front", "toc", "body", "admin-tail"]:
    normalized = _normalize_text(text)
    if not normalized:
        return current
    if normalized == "목 차":
        return "toc"
    if any(re.match(pattern, normalized) for pattern in _BODY_START_PATTERNS):
        return "body"
    if current == "body" and any(re.search(pattern, normalized) for pattern in _ADMIN_TAIL_PATTERNS):
        return "admin-tail"
    return current


def _is_layout_or_noise_table(rows: list[list[str]], state: Literal["front", "toc", "body", "admin-tail"]) -> bool:
    normalized_rows = [[_normalize_text(cell) for cell in row] for row in rows]
    flattened = [cell for row in normalized_rows for cell in row if cell]
    if not flattened:
        return True

    if any(_is_bridge_noise(cell) for cell in flattened):
        return True

    row_count = len(normalized_rows)
    col_count = max((len(row) for row in normalized_rows), default=0)
    nonempty_per_row = [sum(1 for cell in row if cell) for row in normalized_rows]
    blank_slots = sum(1 for row in normalized_rows for cell in row if not cell)
    total_slots = sum(len(row) for row in normalized_rows) or 1
    blank_ratio = blank_slots / total_slots
    joined = " ".join(flattened)

    if max(nonempty_per_row, default=0) <= 1 and len(flattened) <= 2:
        return True

    if len(flattened) == 1 and len(flattened[0]) > 80:
        return True

    if "작성요령" in joined or "제출 시 삭제" in joined or "유의사항" in joined:
        return True

    if col_count >= 8:
        return True
    if row_count >= 10 and blank_ratio >= 0.2:
        return True
    if blank_ratio >= 0.45:
        return True
    if row_count >= 12:
        return True
    if any(re.search(pattern, joined) for pattern in _ADMIN_TAIL_PATTERNS):
        return True

    if state in {"front", "toc", "admin-tail"}:
        if len(flattened) == 1:
            return True
        if max(nonempty_per_row, default=0) <= 1 and row_count <= 2:
            return True

    return False


def _paragraph_overlaps_table(text: str, rows: list[list[str]]) -> bool:
    normalized = _normalize_text(text)
    if not normalized:
        return False
    table_bits = [_normalize_text(cell) for row in rows for cell in row if _normalize_text(cell)]
    if not table_bits:
        return False
    joined = " ".join(table_bits)
    compact_text = re.sub(r"\s+", "", normalized)
    compact_table = re.sub(r"\s+", "", joined)
    return compact_text == compact_table or compact_text in compact_table


def parse_hwpx_to_model(input_path: Path) -> DocumentModel:
    model = DocumentModel()
    saw_paragraph = False
    parser_state: Literal["front", "toc", "body", "admin-tail"] = "front"
    bridge_like_front_matter = False
    toc_mode = False
    seen_paragraphs: set[str] = set()
    seen_tables: set[tuple[tuple[str, ...], ...]] = set()
    recent_table_rows: list[list[str]] | None = None

    with zipfile.ZipFile(input_path) as zf:
        for member in _iter_xml_members(zf):
            try:
                root = ET.fromstring(zf.read(member))
            except ET.ParseError:
                continue

            skip_nodes: set[int] = set()
            for node in root.iter():
                local = _local_name(node.tag).lower()
                attrs = {(_local_name(k).lower()): v for k, v in node.attrib.items()}
                style_hint = " ".join(attrs.values()).lower()
                trace = f"{member}:{local}"

                if id(node) in skip_nodes:
                    continue

                if local == "tbl":
                    if parser_state in {"front", "toc"} and bridge_like_front_matter:
                        # Ignore template/admin tables before body starts for known front-heavy templates.
                        for descendant in node.iter():
                            skip_nodes.add(id(descendant))
                        continue

                    rows = _extract_table_rows(node)
                    signature = _table_signature(rows)
                    if rows and signature not in seen_tables and not _is_layout_or_noise_table(rows, parser_state):
                        seen_tables.add(signature)
                        model.add_table(rows, source_trace=trace)
                        recent_table_rows = rows
                    for descendant in node.iter():
                        skip_nodes.add(id(descendant))
                    continue

                if local in _PARAGRAPH_TAG_CANDIDATES:
                    saw_paragraph = True
                    raw_runs = _extract_inline_runs(node, trace)
                    text = "".join(run.text for run in raw_runs) if raw_runs else _node_text(node)
                    for chunk in _split_glued_paragraph(text):
                        normalized = _normalize_text(chunk)
                        if _is_bridge_noise(normalized):
                            continue
                        if recent_table_rows and _paragraph_overlaps_table(normalized, recent_table_rows):
                            continue
                        if normalized in seen_paragraphs:
                            continue
                        parser_state = _detect_parser_state(normalized, parser_state)
                        if normalized == "목 차":
                            toc_mode = True
                        if any(token in normalized for token in ("글로컬랩", "2025년도", "건국대학교", "작성요령", "제출 시 삭제", "유의사항")):
                            bridge_like_front_matter = True
                        if normalized in {"작성요령", "작성요령 (제출 시 삭제)", "유의사항", "글로컬랩", "건국대학교 첨단재생과학연구원", "설립"}:
                            continue
                        if any(re.match(pattern, normalized) for pattern in _FRONT_COVER_NOISE_PATTERNS):
                            continue
                        if parser_state == "front" and model.blocks and not bridge_like_front_matter:
                            parser_state = "body"
                        chunk_runs = []
                        if len(_split_glued_paragraph(text)) == 1:
                            chunk_runs = raw_runs
                        if parser_state in {"front", "toc"}:
                            if not bridge_like_front_matter:
                                # Conservative behavior for generic documents: preserve front/moc blocks.
                                if normalized.startswith("2025년도"):
                                    model.add_paragraph(normalized, source_trace=trace, runs=chunk_runs)
                                    seen_paragraphs.add(normalized)
                                elif normalized == "목 차":
                                    model.add_heading(normalized, level=1, source_trace=trace, runs=chunk_runs)
                                    seen_paragraphs.add(normalized)
                                elif parser_state == "front" and _should_promote_to_heading(normalized, style_hint, body_started=False):
                                    model.add_heading(normalized, level=1, source_trace=trace, runs=chunk_runs)
                                    seen_paragraphs.add(normalized)
                                continue

                            # Template-heavy front matter: keep only explicit structural headings in front/toc.
                            if normalized == "목 차":
                                model.add_heading(normalized, level=1, source_trace=trace, runs=chunk_runs)
                                seen_paragraphs.add(normalized)
                                continue
                            if _should_promote_to_heading(normalized, style_hint, body_started=False):
                                model.add_heading(normalized, level=1, source_trace=trace, runs=chunk_runs)
                                seen_paragraphs.add(normalized)
                            continue
                        if parser_state == "admin-tail":
                            continue
                        if toc_mode and any(__import__("re").match(pattern, normalized) for pattern in _SECTION_HEADING_PATTERNS):
                            seen_paragraphs.add(normalized)
                            continue
                        seen_paragraphs.add(normalized)
                        if normalized.startswith("(") and not style_hint:
                            model.add_paragraph(normalized, source_trace=trace, runs=chunk_runs)
                            continue
                        if _should_promote_to_heading(normalized, style_hint, body_started=(parser_state == "body")):
                            model.add_heading(normalized, level=1, source_trace=trace, runs=chunk_runs)
                        else:
                            model.add_paragraph(normalized, source_trace=trace, runs=chunk_runs)
                    recent_table_rows = None
                    continue

                if not saw_paragraph and local in _TEXT_TAG_CANDIDATES:
                    text = _node_text(node).strip()
                    for chunk in _split_glued_paragraph(text):
                        normalized = _normalize_text(chunk)
                        if _is_bridge_noise(normalized) or normalized in seen_paragraphs:
                            continue
                        if any(re.match(pattern, normalized) for pattern in _FRONT_COVER_NOISE_PATTERNS):
                            continue
                        seen_paragraphs.add(normalized)
                        if parser_state == "body" and _should_promote_to_heading(normalized, style_hint, body_started=True):
                            model.add_heading(normalized, level=1, source_trace=trace)
                        else:
                            model.add_paragraph(normalized, source_trace=trace)

    return model


def extract_hwpx_text(input_path: Path) -> str:
    return render_markdown(parse_hwpx_to_model(input_path))


def _build_result(
    input_path: Path,
    output_path: Path,
    source_format: str,
    target_format: str,
    route: str | None,
    template: Path | None,
    preserve_original: bool,
) -> ConversionResult:
    chosen_route = route or choose_route(source_format, target_format)
    return ConversionResult(
        input_path=str(input_path),
        output_path=str(output_path),
        source_format=source_format,
        target_format=target_format,
        route=chosen_route,
        preserve_original=preserve_original,
        template=str(template) if template else None,
        input_size_bytes=input_path.stat().st_size,
        output_size_bytes=output_path.stat().st_size,
    )


def run_hwpx_to_md_conversion(
    input_path: Path,
    output_path: Path,
    source_format: str,
    target_format: str,
    route: str | None = None,
    template: Path | None = None,
    preserve_original: bool = True,
) -> ConversionResult:
    if source_format != "hwpx" or target_format != "md":
        raise ValueError("run_hwpx_to_md_conversion only supports hwpx -> md")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    extracted = extract_hwpx_text(input_path)
    output_path.write_text(extracted, encoding="utf-8")

    return _build_result(
        input_path=input_path,
        output_path=output_path,
        source_format=source_format,
        target_format=target_format,
        route=route,
        template=template,
        preserve_original=preserve_original,
    )


def run_hwpx_to_docx_conversion(
    input_path: Path,
    output_path: Path,
    source_format: str,
    target_format: str,
    route: str | None = None,
    template: Path | None = None,
    preserve_original: bool = True,
) -> ConversionResult:
    if source_format != "hwpx" or target_format != "docx":
        raise ValueError("run_hwpx_to_docx_conversion only supports hwpx -> docx")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    model = parse_hwpx_to_model(input_path)
    write_docx(model, output_path)

    return _build_result(
        input_path=input_path,
        output_path=output_path,
        source_format=source_format,
        target_format=target_format,
        route=route,
        template=template,
        preserve_original=preserve_original,
    )


def _write_hwpx_from_model(
    model: DocumentModel,
    input_path: Path,
    output_path: Path,
    source_format: str,
    target_format: str,
    route: str | None,
    template: Path | None,
    preserve_original: bool,
) -> ConversionResult:
    write_hwpx(model, output_path, reference_path=template)
    return _build_result(
        input_path=input_path,
        output_path=output_path,
        source_format=source_format,
        target_format=target_format,
        route=route,
        template=template,
        preserve_original=preserve_original,
    )


def run_md_to_hwpx_conversion(
    input_path: Path,
    output_path: Path,
    source_format: str,
    target_format: str,
    route: str | None = None,
    template: Path | None = None,
    preserve_original: bool = True,
) -> ConversionResult:
    if source_format != "md" or target_format != "hwpx":
        raise ValueError("run_md_to_hwpx_conversion only supports md -> hwpx")

    markdown_text = input_path.read_text(encoding="utf-8", errors="ignore")
    model = parse_markdown(markdown_text)
    return _write_hwpx_from_model(
        model=model,
        input_path=input_path,
        output_path=output_path,
        source_format=source_format,
        target_format=target_format,
        route=route,
        template=template,
        preserve_original=preserve_original,
    )


def run_docx_to_hwpx_conversion(
    input_path: Path,
    output_path: Path,
    source_format: str,
    target_format: str,
    route: str | None = None,
    template: Path | None = None,
    preserve_original: bool = True,
) -> ConversionResult:
    if source_format != "docx" or target_format != "hwpx":
        raise ValueError("run_docx_to_hwpx_conversion only supports docx -> hwpx")

    from docx import Document

    document = Document(input_path)
    model = DocumentModel()
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or "").lower() if para.style else ""
        runs = [
            InlineRun(
                text=run.text,
                bold=bool(run.bold),
                italic=bool(run.italic),
                underline=bool(run.underline),
                source_trace="docx:run",
            )
            for run in para.runs
            if run.text
        ]
        if "heading" in style_name or style_name == "title":
            model.add_heading(text, level=1, source_trace="docx:heading", runs=runs)
        else:
            model.add_paragraph(text, source_trace="docx:paragraph", runs=runs)

    return _write_hwpx_from_model(
        model=model,
        input_path=input_path,
        output_path=output_path,
        source_format=source_format,
        target_format=target_format,
        route=route,
        template=template,
        preserve_original=preserve_original,
    )
