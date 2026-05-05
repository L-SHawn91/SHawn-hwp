"""DOCX writer helpers for SHawn-hwp DocumentModel."""

from __future__ import annotations

from pathlib import Path

from shawn_hwp.model import DocumentModel


def _write_runs(paragraph, block) -> None:
    if block.runs:
        for run_spec in block.runs:
            run = paragraph.add_run(run_spec.text)
            run.bold = run_spec.bold
            run.italic = run_spec.italic
            run.underline = run_spec.underline
    else:
        paragraph.add_run(block.text)


def _apply_table_cell_spans(table, rows: list[list[str]], cell_spans: list[dict[str, int]]) -> None:
    """Apply DocumentModel table span metadata to a python-docx table."""

    if not rows or not cell_spans:
        return
    height = len(rows)
    width = max(len(row) for row in rows)
    for span in cell_spans:
        row = int(span.get("row", 0))
        col = int(span.get("col", 0))
        rowspan = max(int(span.get("rowspan", 1)), 1)
        colspan = max(int(span.get("colspan", 1)), 1)
        end_row = row + rowspan - 1
        end_col = col + colspan - 1
        if row < 0 or col < 0 or end_row >= height or end_col >= width:
            continue
        origin_text = rows[row][col] if col < len(rows[row]) else ""
        merged_cell = table.cell(row, col).merge(table.cell(end_row, end_col))
        merged_cell.text = origin_text


def write_docx(model: DocumentModel, output_path: Path) -> None:
    from docx import Document

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for block in model.blocks:
        if block.kind == "heading":
            level = max(1, min(block.level or 1, 9))
            paragraph = document.add_heading("", level=level)
            _write_runs(paragraph, block)
        elif block.kind == "paragraph":
            paragraph = document.add_paragraph("")
            _write_runs(paragraph, block)
        elif block.kind == "table":
            rows = block.rows
            if not rows:
                continue
            width = max(len(r) for r in rows)
            table = document.add_table(rows=len(rows), cols=width)
            for r_idx, row in enumerate(rows):
                for c_idx in range(width):
                    table.cell(r_idx, c_idx).text = row[c_idx] if c_idx < len(row) else ""
            _apply_table_cell_spans(table, rows, block.cell_spans)
    document.save(output_path)
